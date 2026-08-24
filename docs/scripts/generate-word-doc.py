#!/usr/bin/env python3
"""Generate OCP-4.18-SNO-Disconnected-Upgrade-SOP.docx from the docs package."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "OCP-4.18-SNO-Disconnected-Upgrade-SOP.docx"

SECTIONS = [
    ROOT / "README.md",
    ROOT / "docs/00-overview-constraints.md",
    ROOT / "docs/01-sop-upgrade.md",
    ROOT / "docs/02-pre-during-post-checklist.md",
    ROOT / "docs/03-troubleshooting-guide.md",
    ROOT / "docs/04-second-sno-cluster.md",
    ROOT / "docs/05-failure-cases.md",
    ROOT / "docs/06-best-practices.md",
    ROOT / "docs/07-mindmap-and-diagrams.md",
    ROOT / "docs/08-quick-reference.md",
]


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n\n\n")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(
        "OpenShift SNO Disconnected Upgrade\nStandard Operating Procedure"
    )
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = sub.add_run(
        "Scenario: OCP 4.18.6 → 4.18.52 (ICSP + Quay mirror)\n"
        "Includes: constraints, checklists, troubleshooting,\n"
        "2nd SNO API parity, failure cases, diagrams"
    )
    s.font.size = Pt(12)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = meta.add_run(
        "\nDocument type: Operational SOP / Runbook\n"
        "Derived from lab conversation incident analysis\n"
        "Follow Phase 0 mirror gates before any upgrade"
    )
    m.font.size = Pt(10)
    m.italic = True
    doc.add_page_break()


def add_markdown_file(doc: Document, path: Path) -> None:
    text = path.read_text(encoding="utf-8")
    doc.add_heading(path.name, level=1)
    in_code = False
    code_lines: list[str] = []

    def flush_code() -> None:
        nonlocal code_lines
        if not code_lines:
            return
        p = doc.add_paragraph()
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Consolas"
        run.font.size = Pt(8)
        code_lines = []

    for line in text.splitlines():
        if line.startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("|") and "|" in line[1:]:
            # Keep tables as preformatted text for simplicity/reliability
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = "Consolas"
            run.font.size = Pt(8)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip().startswith(("1. ", "2. ", "3. ", "4. ", "5. ", "6. ", "7. ", "8. ", "9. ")):
            # numbered list rough handling
            content = line.strip().split(". ", 1)[-1]
            doc.add_paragraph(content, style="List Number")
        elif line.strip() == "---":
            doc.add_paragraph("—" * 20)
        elif line.strip() == "":
            continue
        else:
            # strip simple bold markers for readability
            cleaned = line.replace("**", "")
            doc.add_paragraph(cleaned)

    if in_code:
        flush_code()
    doc.add_page_break()


def main() -> None:
    doc = Document()
    set_doc_defaults(doc)
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    add_title_page(doc)

    toc = doc.add_heading("Contents (source files)", level=1)
    for p in SECTIONS:
        doc.add_paragraph(str(p.relative_to(ROOT)), style="List Number")
    doc.add_page_break()

    for path in SECTIONS:
        if not path.exists():
            doc.add_paragraph(f"MISSING: {path}")
            continue
        add_markdown_file(doc, path)

    # Appendix: script names
    doc.add_heading("Appendix — Automation scripts", level=1)
    scripts = sorted((ROOT / "docs/scripts").glob("*.sh"))
    for s in scripts:
        doc.add_paragraph(s.name, style="List Bullet")
        doc.add_paragraph(s.read_text(encoding="utf-8")[:1500] + "\n...")

    doc.add_heading("Appendix — API resource templates", level=1)
    for y in sorted((ROOT / "docs/api-resources").glob("*.yaml")):
        doc.add_heading(y.name, level=2)
        p = doc.add_paragraph()
        run = p.add_run(y.read_text(encoding="utf-8"))
        run.font.name = "Consolas"
        run.font.size = Pt(8)

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
